import os
from io import BytesIO
from pathlib import Path

import pytest

from src.core.document_parser import DocumentParser
from src.exceptions import UnsupportedFileTypeError


FIXTURES_DIR = Path(__file__).parent / "fixtures"


class TestDocumentParser:
    def setup_method(self):
        self.parser = DocumentParser()

    def test_parse_txt_returns_text(self):
        """TXT files should return raw text content."""
        content = b"This is a test document with some text."
        result = self.parser.parse(content, "test.txt")
        assert result.text == "This is a test document with some text."
        assert result.file_type == "txt"
        assert result.filename == "test.txt"

    def test_parse_txt_fixture(self):
        """TXT fixture should parse correctly."""
        content = (FIXTURES_DIR / "sample_rate_confirmation.txt").read_bytes()
        result = self.parser.parse(content, "sample_rate_confirmation.txt")
        assert "RATE CONFIRMATION" in result.text
        assert "Swift Transportation" in result.text
        assert "$2,450.00" in result.text
        assert result.file_type == "txt"

    def test_parse_txt_multiline(self):
        """Multi-line TXT should preserve line breaks."""
        content = b"Line 1\nLine 2\nLine 3"
        result = self.parser.parse(content, "multi.txt")
        assert "Line 1" in result.text
        assert "Line 2" in result.text
        assert "Line 3" in result.text

    def test_parse_empty_file_returns_empty(self):
        """Empty files should return empty string."""
        result = self.parser.parse(b"", "empty.txt")
        assert result.text == ""

    def test_parse_unsupported_raises_error(self):
        """Unsupported file types should raise UnsupportedFileTypeError."""
        with pytest.raises(UnsupportedFileTypeError):
            self.parser.parse(b"data", "file.exe")

    def test_parse_unsupported_xlsx(self):
        """XLSX should raise UnsupportedFileTypeError."""
        with pytest.raises(UnsupportedFileTypeError):
            self.parser.parse(b"data", "file.xlsx")

    def test_parse_pdf_extracts_text(self):
        """PDF files should extract readable text."""
        # Create a minimal PDF for testing
        pdf_path = FIXTURES_DIR / "sample_rate_confirmation.pdf"
        if pdf_path.exists():
            content = pdf_path.read_bytes()
            result = self.parser.parse(content, "test.pdf")
            assert result.file_type == "pdf"
            assert len(result.text) > 0

    def test_parse_docx_extracts_text(self):
        """DOCX files should extract paragraph text."""
        # Create a simple DOCX in memory for testing
        from docx import Document
        doc = Document()
        doc.add_paragraph("Test paragraph one")
        doc.add_paragraph("Test paragraph two")
        buffer = BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        result = self.parser.parse(content, "test.docx")
        assert "Test paragraph one" in result.text
        assert "Test paragraph two" in result.text
        assert result.file_type == "docx"

    def test_parse_returns_parsed_document_type(self):
        """Parser should return ParsedDocument dataclass."""
        result = self.parser.parse(b"Hello", "test.txt")
        assert hasattr(result, "text")
        assert hasattr(result, "filename")
        assert hasattr(result, "file_type")
        assert hasattr(result, "num_pages")
        assert hasattr(result, "tables")

    def test_parse_pdf_returns_page_count(self):
        """PDF parsing should return correct page count."""
        pdf_path = FIXTURES_DIR / "sample_rate_confirmation.pdf"
        if pdf_path.exists():
            content = pdf_path.read_bytes()
            result = self.parser.parse(content, "test.pdf")
            assert result.num_pages >= 1

    def test_file_type_detection_case_insensitive(self):
        """File type detection should be case-insensitive."""
        result = self.parser.parse(b"test", "FILE.TXT")
        assert result.file_type == "txt"

    def test_parse_txt_strips_bom(self):
        """TXT parser should handle UTF-8 BOM."""
        content = b"\xef\xbb\xbfHello World"
        result = self.parser.parse(content, "bom.txt")
        assert result.text == "Hello World"
