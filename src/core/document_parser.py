"""Multi-format document text extraction.

Supports PDF (via pdfplumber with PyPDF2 fallback), DOCX, and plain text.
Designed for semi-structured logistics documents where preserving
whitespace and table structure matters for downstream chunking.
"""

import io
from typing import List

from src.exceptions import DocumentParsingError, UnsupportedFileTypeError
from src.models.document import ParsedDocument
from src.util.logging_setup import get_logger

logger = get_logger(__name__)

SUPPORTED_TYPES = {"pdf", "docx", "txt"}


class DocumentParser:
    """Extracts raw text from uploaded documents.

    Each format handler returns a ``ParsedDocument`` containing the full
    extracted text, page count, and any detected tables (PDF only).
    """

    def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        """Parse a document and return extracted text.

        Args:
            file_bytes: Raw file content.
            filename: Original filename (used to determine type).

        Returns:
            ParsedDocument with text, metadata, and tables.

        Raises:
            UnsupportedFileTypeError: If the file extension is not in SUPPORTED_TYPES.
            DocumentParsingError: If text extraction fails.
        """
        suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if suffix not in SUPPORTED_TYPES:
            raise UnsupportedFileTypeError(suffix)

        logger.info("Parsing document: %s (%s, %d bytes)", filename, suffix, len(file_bytes))

        if suffix == "txt":
            return self._parse_txt(file_bytes, filename)
        elif suffix == "pdf":
            return self._parse_pdf(file_bytes, filename)
        else:  # docx — guaranteed by SUPPORTED_TYPES check
            return self._parse_docx(file_bytes, filename)

    def _parse_txt(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        try:
            text = file_bytes.decode("utf-8-sig").strip()
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1").strip()
        return ParsedDocument(text=text, filename=filename, file_type="txt", num_pages=1)

    def _parse_pdf(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        try:
            import pdfplumber
        except ImportError:
            raise DocumentParsingError(filename, "pdfplumber not installed — run: pip install pdfplumber")

        tables: List[List[List[str]]] = []
        pages_text = []

        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                num_pages = len(pdf.pages)
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    pages_text.append(text)
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)
        except Exception as exc:
            raise DocumentParsingError(filename, str(exc)) from exc

        full_text = "\n".join(pages_text).strip()

        # Fallback to PyPDF2 if pdfplumber yields no text (e.g., encrypted PDF)
        if not full_text:
            logger.warning("pdfplumber returned no text for %s, falling back to PyPDF2", filename)
            full_text = self._parse_pdf_fallback(file_bytes, filename)

        logger.info("Parsed PDF: %d pages, %d chars, %d tables", num_pages, len(full_text), len(tables))
        return ParsedDocument(
            text=full_text, filename=filename, file_type="pdf",
            num_pages=num_pages, tables=tables, pages=pages_text,
        )

    def _parse_pdf_fallback(self, file_bytes: bytes, filename: str) -> str:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise DocumentParsingError(filename, "PyPDF2 not installed — run: pip install PyPDF2")

        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()

    def _parse_docx(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError:
            raise DocumentParsingError(filename, "python-docx not installed — run: pip install python-docx")

        try:
            doc = Document(io.BytesIO(file_bytes))
        except Exception as exc:
            raise DocumentParsingError(filename, str(exc)) from exc

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        return ParsedDocument(text=text, filename=filename, file_type="docx", num_pages=1)
